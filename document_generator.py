import json
import re
import os
from typing import Dict, List, Optional, Generator
from openai import OpenAI
from prompt_manager import PromptManager
from datetime import datetime
import time
from docx import Document
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement

class DocumentCleaner:
    """Handles document cleaning and formatting operations"""
    
    @staticmethod
    def set_cell_shading(cell, color):
        """Set cell background color"""
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    @staticmethod
    def is_table_row(line):
        """Check if a line contains tab-separated table data"""
        return len(re.split(r'\s{2,}', line.strip())) > 1
    
    @staticmethod
    def insert_page_break_after(element):
        """Insert page break after an element"""
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        r.append(br)
        p.append(r)
        element.addnext(p)
    
    @staticmethod
    def convert_text_to_formatted_docx(text_content: str, output_path: str) -> str:
        """Convert text content to formatted DOCX with tables"""
        # Create a new document
        doc = Document()
        
        # Split content into paragraphs
        lines = text_content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if line.startswith("Table Title:"):
                # Collect table data
                paragraphs_to_skip = 0
                table_rows = []
                
                # Add title
                title_paragraph = doc.add_paragraph(line)
                title_paragraph.style = 'Heading 3'
                
                i += 1
                paragraphs_to_skip += 1
                
                # Skip empty lines
                while i < len(lines) and not lines[i].strip():
                    i += 1
                    paragraphs_to_skip += 1
                
                # Header
                if i < len(lines) and DocumentCleaner.is_table_row(lines[i]):
                    header = re.split(r'\s{2,}', lines[i].strip())
                    table_rows.append(header)
                    i += 1
                    paragraphs_to_skip += 1
                else:
                    # Not a valid table, continue
                    continue
                
                # Data rows
                while i < len(lines):
                    line = lines[i].strip()
                    if not DocumentCleaner.is_table_row(line):
                        break
                    row = re.split(r'\s{2,}', line)
                    table_rows.append(row)
                    i += 1
                    paragraphs_to_skip += 1
                
                # Create table
                if table_rows:
                    cols = len(table_rows[0])
                    table = doc.add_table(rows=1, cols=cols)
                    table.style = "Table Grid"
                    
                    # Header formatting
                    for c, val in enumerate(table_rows[0]):
                        cell = table.rows[0].cells[c]
                        cell.text = val
                        DocumentCleaner.set_cell_shading(cell, "D9EAD3")
                        cell.paragraphs[0].runs[0].bold = True
                    
                    # Body formatting
                    for row_data in table_rows[1:]:
                        cells = table.add_row().cells
                        for c, val in enumerate(row_data):
                            cells[c].text = val
                            if c == 0:
                                DocumentCleaner.set_cell_shading(cells[c], "D9EAD3")
                                cells[c].paragraphs[0].runs[0].bold = True
                    
                    # Add page break after table
                    if i < len(lines) - 1:
                        doc.add_page_break()
                
                continue
            
            elif line.startswith('# '):
                # Main title
                doc.add_heading(line[2:], level=0)
            
            elif line.startswith('## '):
                # Section heading
                doc.add_heading(line[3:], level=1)
            
            elif line.startswith('### '):
                # Subsection heading
                doc.add_heading(line[4:], level=2)
            
            elif line.startswith('#### '):
                # Sub-subsection heading
                doc.add_heading(line[5:], level=3)
            
            elif line.strip() == '---' or line.strip() == '***':
                # Horizontal line
                p = doc.add_paragraph()
                p.add_run().add_break()
            
            elif line.strip():
                # Regular paragraph
                doc.add_paragraph(line)
            
            i += 1
        
        # Save document
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def clean_and_format_markdown(document_text: str) -> str:
        """Clean and format markdown document before conversion"""
        lines = document_text.split('\n')
        cleaned_lines = []
        
        # Remove excessive blank lines
        blank_line_count = 0
        for line in lines:
            if line.strip() == '':
                blank_line_count += 1
                if blank_line_count <= 2:  # Keep max 2 consecutive blank lines
                    cleaned_lines.append(line)
            else:
                blank_line_count = 0
                cleaned_lines.append(line)
        
        # Ensure proper table spacing
        formatted_lines = []
        in_table = False
        
        for i, line in enumerate(cleaned_lines):
            current_line = line
            
            # Check for table title
            if line.strip().startswith('Table Title:'):
                # Add 2 blank lines before table title
                if i > 0 and cleaned_lines[i-1].strip() != '':
                    formatted_lines.append('')
                    formatted_lines.append('')
                formatted_lines.append(current_line)
                in_table = True
                continue
            
            # Check for table data
            if in_table:
                if DocumentCleaner.is_table_row(current_line):
                    formatted_lines.append(current_line)
                else:
                    # Table ended, add 3 blank lines
                    formatted_lines.append('')
                    formatted_lines.append('')
                    formatted_lines.append('')
                    formatted_lines.append(current_line)
                    in_table = False
            else:
                formatted_lines.append(current_line)
        
        # Add final spacing if still in table
        if in_table:
            formatted_lines.append('')
            formatted_lines.append('')
            formatted_lines.append('')
        
        return '\n'.join(formatted_lines)

class DocumentGenerator:
    def __init__(self, client: OpenAI, prompt_manager: PromptManager):
        self.client = client
        self.prompt_manager = prompt_manager
        self.max_tokens_per_request = 4000
        self.document_cleaner = DocumentCleaner()
    
    def generate_document(self, project: Dict, document_type: str, 
                         detail_level: str = "Comprehensive",
                         include_diagrams: bool = False,
                         format_tables: bool = False) -> str:
        """Generate document based on document type and detail level"""
        project_name = project["product_name"]
        project_details = json.dumps(project["sections"], indent=2)
        
        print(f"Generating {document_type} at {detail_level} level...")
        
        document = self._generate_comprehensive_document(
                project, document_type, project_name, project_details, 
                include_diagrams, format_tables
            )
        
        # Clean and format the document
        cleaned_document = self.document_cleaner.clean_and_format_markdown(document)
        
        # Generate DOCX file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"{document_type.replace(' ', '_')}_{detail_level}_{timestamp}.docx"
        output_dir = "generated_documents"
        os.makedirs(output_dir, exist_ok=True)
        docx_path = os.path.join(output_dir, docx_filename)
        
        self.document_cleaner.convert_text_to_formatted_docx(cleaned_document, docx_path)
        print(f"Generated DOCX file: {docx_path}")
        
        return cleaned_document
    
    def _generate_comprehensive_document(self, project: Dict, document_type: str, 
                                        project_name: str, project_details: str,
                                        include_diagrams: bool, format_tables: bool) -> str:
        """Generate comprehensive document by sections"""
        sections = self._get_document_sections(document_type)
        
        print(f"Generating {len(sections)} sections for {document_type}...")
        
        document = f"""# {document_type}
## Table of Contents
"""        
        for i, section in enumerate(sections, 1):
            section_name = section["name"]
            document += f"{section_name}\n"
        
        document += f"{len(sections) + 1}. Appendices\n"
        document += "\n---\n\n"
        
        for i, section in enumerate(sections, 1):
            section_name = section["name"]
            section_instructions = section["instructions"]
            
            print(f"Generating section {i}/{len(sections)}: {section_name}")
            
            document += f"# {section_name}\n\n"
            
            # Check if this section needs tables
            needs_tables = self._section_needs_tables(document_type, section_name)
            
            section_content = self._generate_document_section(
                project, project_name, project_details, section_name, 
                section_instructions, include_diagrams, document_type, needs_tables
            )
            
            # Remove the first line if it's a heading (to avoid duplication)
            lines = section_content.splitlines()
            if lines and lines[0].startswith('#'):
                section_content = '\n'.join(lines[1:])
            
            document += section_content
            document += "\n\n---\n\n"
            
            if i < len(sections):
                time.sleep(1)
        
        document += self._generate_comprehensive_appendices(project, document_type)
        
        return document
    
    def _section_needs_tables(self, document_type: str, section_name: str) -> bool:
        """Check if a section needs tables based on the specification"""
        table_sections = {
            "Business Requirements Document (BRD)": [
                "7. Cost-Benefit Analysis"
            ],
            "Functional Requirements Document (FRD)": [
                "2. Functional Requirements",
                "4. Data & Integration Requirements"
            ],
            "Non-Functional Requirements (NFR)": [
                "1. Performance",
                "2. Security",
                "3. Availability & Reliability"
            ],
            "User Stories & Acceptance Criteria": [
                "1. Epics & Themes",
                "2. User Stories",
                "3. Acceptance Criteria"
            ],
            "Stakeholder Analysis Matrix": [
                "2. Power vs Interest"
            ],
            "SWOT Analysis Report": [
                "1. Internal Factors",
                "2. External Factors",
                "3. Strategic Analysis"
            ],
            "Process Flow Diagrams (BPMN)": []  # No tables needed for BPMN
        }
        
        return section_name in table_sections.get(document_type, [])
    
    def _generate_document_section(self, project: Dict, project_name: str,
                               project_details: str, section_name: str,
                               section_instructions: str, include_diagrams: bool,
                               document_type: str, needs_tables: bool = False) -> str:
        """Generate a single document section with plain text tables"""
        
        # Base prompt
        base_prompt = f"""
PROJECT: {project_name}
SECTION: {section_name}
PROJECT DETAILS: {project_details}
INSTRUCTIONS: {section_instructions}
"""
        
        # Adjust word count based on document type
        if document_type in ["Business Requirements Document (BRD)", 
                           "Functional Requirements Document (FRD)",
                           "SWOT Analysis Report"]:
            base_prompt += "\nGenerate 1000-1200 words of professional, well-structured content.\n"
        else:
            base_prompt += "\nGenerate 800-1000 words of professional, well-structured content.\n"
        
        # Add table formatting instructions if needed
        if needs_tables:
            table_instructions = self._get_table_instructions(document_type, section_name)
            base_prompt += f"\n{table_instructions}"
        
        # Add diagram instructions if applicable
        if include_diagrams:
            diagram_hint = self._get_diagram_hint(document_type, section_name)
            if diagram_hint:
                base_prompt += f"\n{diagram_hint}"
        
        # Final prompt
        prompt = base_prompt.strip()
        
        try:
            # Select model based on document complexity
            if document_type in ["Business Requirements Document (BRD)", 
                               "Functional Requirements Document (FRD)",
                               "SWOT Analysis Report"]:
                model = "gpt-4"
            else:
                model = "gpt-4"
            
            response = self.client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": self._get_system_prompt(document_type, section_name, needs_tables)},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
                max_tokens=2500  # Increased for longer content
            )
            
            content = response.choices[0].message.content.strip()
            
            # Clean up table formatting
            if needs_tables:
                content = self._clean_table_formatting(content)
            
            word_count = len(content.split())
            print(f"  Generated {word_count} words for {section_name}")
            
            return content
            
        except Exception as e:
            print(f"Error generating section {section_name}: {str(e)}")
            return f"*Error generating content for {section_name}: {str(e)}*"
    
    def _get_system_prompt(self, document_type: str, section_name: str, needs_tables: bool) -> str:
        """Get system prompt based on section type"""
        base_prompt = "You are a senior Business Analyst and technical documentation expert."
        
        if needs_tables:
            base_prompt += """ CRITICAL TABLE FORMATTING:
1. ALWAYS start each table with "Table Title: [Table Name]" on its own line
2. Leave 2 blank lines before the table title
3. Use TAB characters (press Tab key) between columns - NOT spaces
4. Each table row must be on a single line with TABs separating columns
5. Column headers should be on the first row after the title, separated by TABs
6. Leave 3 blank lines after the table ends
7. Never use pipe characters (|) or ASCII art tables
8. Make sure tables are properly formatted for MS Word

EXAMPLE:
Table Title: Sample Table
Column1    Column2    Column3
Data1      Data2      Data3
More1      More2      More3



"""
        
        # Add document-specific instructions
        if document_type == "Business Requirements Document (BRD)":
            base_prompt += "\nFocus on business needs, objectives, and value proposition."
        elif document_type == "Functional Requirements Document (FRD)":
            base_prompt += "\nFocus on system behavior, user interactions, and functional specifications."
        elif document_type == "SWOT Analysis Report":
            base_prompt += "\nFocus on strategic analysis, competitive positioning, and actionable insights."
        
        return base_prompt
    
    def _get_table_instructions(self, document_type: str, section_name: str) -> str:
        """Get specific table instructions for each section"""
        
        instructions = """
CRITICAL TABLE FORMATTING:
- Start each table with "Table Title: [Table Name]" on its own line
- Leave 2 blank lines before the table title
- Use TAB characters (press Tab key) between columns
- Leave 3 blank lines after the table ends
- Never use pipe characters (|) or ASCII art
"""
        
        if document_type == "Business Requirements Document (BRD)":
            if section_name == "7. Cost-Benefit Analysis":
                instructions += """
REQUIRED TABLES:
Table Title: Cost Breakdown
COST CATEGORY    YEAR 1    YEAR 2    YEAR 3    TOTAL
Development    $XX,XXX    $X,XXX    $X,XXX    $XX,XXX

Include 6-8 cost categories covering development, infrastructure, maintenance, and operational costs.

Table Title: Benefit Analysis
BENEFIT TYPE    YEAR 1    YEAR 2    YEAR 3    TOTAL
Cost Savings    $X,XXX    $XX,XXX    $XX,XXX    $XX,XXX

Include 6-8 benefit types covering cost savings, revenue increase, efficiency gains, and strategic benefits.

Table Title: Financial Metrics
METRIC    VALUE    INTERPRETATION
ROI    XX%    [Interpretation]

Include 4-6 key financial metrics with calculations and business interpretations.
"""
        
        elif document_type == "Functional Requirements Document (FRD)":
            if section_name == "2. Functional Requirements":
                instructions += """
REQUIRED TABLES:
Table Title: Functional Requirements by Module

MODULE    REQUIREMENT ID    DESCRIPTION    PRIORITY
Authentication    FR-001    [Description]    High
User Management    FR-002    [Description]    Medium

Include 12-15 functional requirements across 4-6 system modules.

Table Title: Requirements Priority Matrix
REQUIREMENT    BUSINESS VALUE    TECHNICAL COMPLEXITY    PRIORITY
FR-001    High    Medium    High

Include 8-10 requirements in the priority matrix with clear business value and complexity assessments.
"""
            elif section_name == "4. Data & Integration Requirements":
                instructions += """
REQUIRED TABLE:
Table Title: Data Requirements

DATA ELEMENT    TYPE    FORMAT    VALIDATION RULES
Username    String    alphanumeric    Unique, no spaces

Include 8-10 key data elements with their specifications and validation rules.
"""
        
        elif document_type == "Non-Functional Requirements (NFR)":
            if section_name == "1. Performance":
                instructions += """
REQUIRED TABLES:
Table Title: Performance Requirements

METRIC    REQUIREMENT    ACCEPTANCE CRITERIA
Response Time    < 2 seconds    95% of requests
Throughput    1000 tps    Under peak load

Include 6-8 performance metrics covering response time, throughput, latency, and resource utilization.

Table Title: Performance Benchmarks
SCENARIO    EXPECTED LOAD    TARGET RESPONSE TIME
Normal usage    100 users    < 1 second

Include 4-6 performance scenarios with expected loads and target performance.
"""
            elif section_name == "2. Security":
                instructions += """
REQUIRED TABLES:
Table Title: Security Requirements

REQUIREMENT TYPE    SPECIFICATION    COMPLIANCE
Authentication    Multi-factor    ISO 27001
Encryption    AES-256    GDPR

Include 8-10 security requirements covering authentication, authorization, encryption, and compliance.

Table Title: Access Control Matrix
ROLE    PERMISSIONS    RESTRICTIONS
Admin    Full access    None
User    Read/write own data    No admin functions

Include 4-6 user roles with their specific permissions and restrictions.
"""
            elif section_name == "3. Availability & Reliability":
                instructions += """
REQUIRED TABLE:
Table Title: Availability Requirements

METRIC    TARGET    DESCRIPTION
Uptime    99.9%    Annual availability
MTBF    1000 hours    Mean time between failures

Include 4-6 availability and reliability metrics with clear targets and descriptions.
"""
        
        elif document_type == "User Stories & Acceptance Criteria":
            if section_name == "1. Epics & Themes":
                instructions += """
REQUIRED TABLE:
Table Title: Epics and Themes

EPIC ID    THEME    DESCRIPTION    BUSINESS VALUE
EPIC-001    User Management    Manage user accounts    High

Include 4-6 epics/themes that group related user stories.
"""
            elif section_name == "2. User Stories":
                instructions += """
REQUIRED TABLE:
Table Title: User Stories with Priority

STORY ID    DESCRIPTION    PRIORITY    STORY POINTS
US-001    As a user, I want to login    High    3
US-002    As a user, I want to reset password    Medium    2

Include 10-12 user stories with priority levels and estimated story points.
"""
            elif section_name == "3. Acceptance Criteria":
                instructions += """
REQUIRED TABLE:
Table Title: Acceptance Criteria Matrix

STORY ID    CRITERIA ID    DESCRIPTION    STATUS
US-001    AC-001    User can login with valid credentials    Defined
US-001    AC-002    System displays error for invalid credentials    Defined

Include 3-5 acceptance criteria for each of 4-6 key user stories.
"""
        
        elif document_type == "Stakeholder Analysis Matrix":
            if section_name == "2. Power vs Interest":
                instructions += """
REQUIRED TABLE:
Table Title: Power vs Interest Analysis

STAKEHOLDER    POWER    INTEREST    STRATEGY
[Name]    High    High    Manage closely
[Name]    Low    Low    Monitor

Include 8-10 stakeholders analyzed for their power and interest levels with management strategies.
"""
        
        elif document_type == "SWOT Analysis Report":
            if section_name == "1. Internal Factors":
                instructions += """
REQUIRED TABLES:
Table Title: Strengths Analysis

STRENGTH    DESCRIPTION    IMPACT    EVIDENCE
[Strength]    [Description]    High    [Evidence]

Include 6-8 key strengths with descriptions, impact levels, and supporting evidence.

Table Title: Weaknesses Analysis

WEAKNESS    DESCRIPTION    IMPACT    MITIGATION
[Weakness]    [Description]    High    [Strategy]

Include 6-8 weaknesses with descriptions, impact levels, and mitigation strategies.
"""
            elif section_name == "2. External Factors":
                instructions += """
REQUIRED TABLES:
Table Title: Opportunities Analysis

OPPORTUNITY    DESCRIPTION    PROBABILITY    IMPACT
[Opportunity]    [Description]    High    High

Include 6-8 external opportunities with probability and impact assessments.

Table Title: Threats Analysis

THREAT    DESCRIPTION    PROBABILITY    IMPACT
[Threat]    [Description]    High    High

Include 6-8 external threats with probability and impact assessments.
"""
            elif section_name == "3. Strategic Analysis":
                instructions += """
REQUIRED TABLE:
Table Title: SWOT Matrix

FACTOR    STRENGTHS    WEAKNESSES    OPPORTUNITIES    THREATS
Internal    List strengths    List weaknesses    -    -
External    -    -    List opportunities    List threats

Summarize 3-4 key items in each SWOT category for strategic analysis.
"""
        
        return instructions
    
    def _clean_table_formatting(self, content: str) -> str:
        """Clean up table formatting to ensure proper tab separation"""
        lines = content.split('\n')
        cleaned_lines = []
        
        in_table = False
        table_lines = []
        
        for line in lines:
            # Check if this is a table title
            if line.strip().startswith('Table Title:') or line.strip().startswith('Table:'):
                if in_table and table_lines:
                    # Close previous table
                    cleaned_lines.extend(table_lines)
                    cleaned_lines.append('')  # Add 3 blank lines after table
                    cleaned_lines.append('')
                    cleaned_lines.append('')
                    table_lines = []
                
                # Start new table
                in_table = True
                if 'Table:' in line:
                    line = line.replace('Table:', 'Table Title:')
                table_lines.append('')  # Add 2 blank lines before table
                table_lines.append('')
                table_lines.append(line.strip())
                continue
            
            # If we're in a table and this line has table data
            if in_table:
                if '|' in line:
                    # Convert pipe-separated to tab-separated
                    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                    if cells:
                        # Skip separator rows
                        if not any('---' in cell or '--' in cell for cell in cells):
                            tab_separated = '\t'.join(cells)
                            table_lines.append(tab_separated)
                elif line.strip() == '':
                    # Empty line might end table
                    if table_lines:
                        cleaned_lines.extend(table_lines)
                        cleaned_lines.append('')  # Add 3 blank lines after table
                        cleaned_lines.append('')
                        cleaned_lines.append('')
                    in_table = False
                    table_lines = []
                    cleaned_lines.append(line)
                elif '+---' in line or '+--' in line:
                    # Skip ASCII borders
                    continue
                elif '\t' in line or ('  ' in line and len(line.strip().split()) > 2):
                    # Already tab-separated or appears to be table data
                    table_lines.append(line)
                else:
                    # Not table data, end table
                    if table_lines:
                        cleaned_lines.extend(table_lines)
                        cleaned_lines.append('')  # Add 3 blank lines after table
                        cleaned_lines.append('')
                        cleaned_lines.append('')
                    in_table = False
                    table_lines = []
                    cleaned_lines.append(line)
            else:
                # Not in table, just add the line
                cleaned_lines.append(line)
        
        # Handle any remaining table data
        if table_lines:
            cleaned_lines.extend(table_lines)
            cleaned_lines.append('')  # Add 3 blank lines after table
            cleaned_lines.append('')
            cleaned_lines.append('')
        
        return '\n'.join(cleaned_lines)
    
    def _get_diagram_hint(self, document_type: str, section_name: str) -> str:
        """Get diagram hint for specific sections"""
        diagram_instructions = "\nInclude a Mermaid flowchart diagram to illustrate key concepts."
        
        if document_type == "Business Requirements Document (BRD)":
            if section_name in ["1. Executive Summary", "3. Project Scope"]:
                return diagram_instructions + "\nCreate a high-level system context diagram showing main components and interactions."
        
        elif document_type == "Functional Requirements Document (FRD)":
            if section_name in ["1. Introduction & System Context", "2. Functional Requirements"]:
                return diagram_instructions + "\nCreate process flow diagrams for key user interactions."
        
        elif document_type == "Non-Functional Requirements (NFR)":
            if section_name in ["1. Performance", "2. Security"]:
                return diagram_instructions + "\nCreate diagrams showing performance/scalability or security layers."
        
        elif document_type == "Process Flow Diagrams (BPMN)":
            if section_name == "1. As-Is Process":
                return "Create detailed BPMN-style Mermaid flowchart for current process with swimlanes."
            elif section_name == "2. To-Be Process":
                return "Create optimized BPMN-style Mermaid flowchart with improvements highlighted."
            elif section_name == "3. Decision Points & Actors":
                return "Create swimlane diagram showing actors, decisions, and responsibilities."
        
        elif document_type == "SWOT Analysis Report":
            if section_name == "3. Strategic Analysis":
                return diagram_instructions + "\nCreate a SWOT matrix visualization."
        
        return ""
    
    def _get_document_sections(self, document_type: str) -> List[Dict]:
        """Get sections for comprehensive documents"""
        sections_map = {
            "Business Requirements Document (BRD)": [
                {"name": "1. Executive Summary", "instructions": "Provide high-level overview, business need, and value proposition."},
                {"name": "2. Project Objectives", "instructions": "Define SMART objectives, success criteria, and business goals."},
                {"name": "3. Project Scope", "instructions": "Detail in-scope/out-of-scope items, boundaries, and deliverables."},
                {"name": "4. Business Requirements", "instructions": "List functional and non-functional business needs with priorities."},
                {"name": "5. Key Stakeholders", "instructions": "Identify stakeholders, roles, and communication needs."},
                {"name": "6. Project Constraints", "instructions": "Describe time, budget, technical, and resource constraints."},
                {"name": "7. Cost-Benefit Analysis", "instructions": "Analyze costs, benefits, ROI, and financial justification."},
                {"name": "8. Conclusion", "instructions": "Summarize key points and next steps."}
            ],
            "Functional Requirements Document (FRD)": [
                {"name": "1. Introduction & System Context", "instructions": "Describe system purpose, scope, and user base."},
                {"name": "2. Functional Requirements", "instructions": "Detail system functions, user capabilities, and workflows."},
                {"name": "3. User Stories & Use Cases", "instructions": "Provide user stories, scenarios, and interaction flows."},
                {"name": "4. Data & Integration Requirements", "instructions": "Define data models, storage, and system integrations."},
                {"name": "5. Conclusion", "instructions": "Summarize functional capabilities and implementation approach."}
            ],
            "Non-Functional Requirements (NFR)": [
                {"name": "1. Performance", "instructions": "Specify response times, throughput, scalability, and load handling."},
                {"name": "2. Security", "instructions": "Define authentication, authorization, encryption, and compliance requirements."},
                {"name": "3. Availability & Reliability", "instructions": "Specify uptime, fault tolerance, backup, and recovery requirements."},
                {"name": "4. Usability", "instructions": "Define user interface standards, accessibility, and user experience requirements."},
                {"name": "5. Maintainability", "instructions": "Specify code standards, documentation, and support requirements."}
            ],
            "User Stories & Acceptance Criteria": [
                {"name": "1. Epics & Themes", "instructions": "Define high-level business themes and epic stories."},
                {"name": "2. User Stories", "instructions": "Write detailed user stories with roles, goals, and benefits."},
                {"name": "3. Acceptance Criteria", "instructions": "Define specific, testable acceptance criteria for each story."},
                {"name": "4. Prioritization", "instructions": "Prioritize stories based on value, effort, and dependencies."}
            ],
            "Stakeholder Analysis Matrix": [
                {"name": "1. Introduction", "instructions": "Define purpose and scope of stakeholder analysis."},
                {"name": "2. Power vs Interest", "instructions": "Analyze stakeholder influence and interest levels."},
                {"name": "3. Engagement Strategy", "instructions": "Define communication and engagement approaches for each stakeholder."},
                {"name": "4. Expectations Management", "instructions": "Document stakeholder expectations and concerns."}
            ],
            "SWOT Analysis Report": [
                {"name": "1. Internal Factors", "instructions": "Analyze organizational strengths and weaknesses."},
                {"name": "2. External Factors", "instructions": "Analyze market opportunities and threats."},
                {"name": "3. Strategic Analysis", "instructions": "Cross-reference SWOT factors and derive strategic insights."},
                {"name": "4. Action Plan", "instructions": "Develop actionable strategies based on SWOT analysis."}
            ],
            "Process Flow Diagrams (BPMN)": [
                {"name": "1. As-Is Process", "instructions": "Document current business processes with pain points and inefficiencies."},
                {"name": "2. To-Be Process", "instructions": "Design optimized future processes with improvements and automations."},
                {"name": "3. Decision Points & Actors", "instructions": "Map decision logic, roles, and responsibilities."},
                {"name": "4. Implementation Roadmap", "instructions": "Plan process transition and change management."}
            ]
        }
        
        default_sections = [
            {"name": "1. Executive Summary", "instructions": f"Generate executive summary for {document_type}"},
            {"name": "2. Detailed Analysis", "instructions": f"Generate detailed analysis for {document_type}"},
            {"name": "3. Implementation Plan", "instructions": f"Generate implementation plan for {document_type}"}
        ]
        
        return sections_map.get(document_type, default_sections)
    
    def _generate_comprehensive_appendices(self, project: Dict, document_type: str) -> str:
        """Generate comprehensive appendices"""
        sections = project.get("sections", {})
        
        appendices = """## Appendices

### Appendix A: Project Reference

"""
        
        for section_name, section_content in sections.items():
            if isinstance(section_content, str) and section_content.strip():
                appendices += f"#### {section_name}\n\n"
                appendices += f"{section_content[:500]}...\n\n"
        
        appendices += "\n### Appendix B: Glossary\n\n"
        
        glossaries = {
            "BRD": "Business Requirements Document",
            "FRD": "Functional Requirements Document",
            "NFR": "Non-Functional Requirements",
            "KPI": "Key Performance Indicator",
            "SLA": "Service Level Agreement",
            "RACI": "Responsible, Accountable, Consulted, Informed",
            "API": "Application Programming Interface",
            "CI/CD": "Continuous Integration/Continuous Deployment",
            "ROI": "Return on Investment",
            "TCO": "Total Cost of Ownership",
            "BPMN": "Business Process Model and Notation",
            "SME": "Subject Matter Expert",
            "MVP": "Minimum Viable Product"
        }
        
        for term, definition in glossaries.items():
            appendices += f"- **{term}:** {definition}\n"
        
        appendices += f"""
### Appendix C: Revision History

| Version | Date | Author | Changes | Status |
|---------|------|--------|---------|--------|
| 1.0 | {datetime.now().strftime('%Y-%m-%d')} | AI Business Analyst | Initial Draft | Draft |
| 1.1 | {datetime.now().strftime('%Y-%m-%d')} | AI Business Analyst | Added diagrams and tables | Review |
"""
        
        return appendices