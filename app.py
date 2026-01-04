import streamlit as st
import json
import os
import time
import io
from datetime import datetime
import base64
import traceback
import sys
from project_selector import ProjectSelector
from ba_agent import BusinessAnalystAgent

# Clear ALL caches
try:
    st.cache_resource.clear()
    st.cache_data.clear()
except Exception as e:
    st.warning(f"Could not clear caches: {str(e)}")

# Page configuration
try:
    st.set_page_config(
        page_title="Business Analyst Agent",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="expanded"
    )
except Exception as e:
    st.warning(f"Page configuration issue: {str(e)}")

# Custom CSS for better display
st.markdown("""
<style>
    /* Error Message Styling */
    .error-container {
        background-color: #f8d7da;
        border: 2px solid #f5c6cb;
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
        color: #721c24;
    }
    
    .error-title {
        font-weight: bold;
        font-size: 1.1rem;
        margin-bottom: 8px;
        display: flex;
        align-items: center;
        gap: 8px;
    }
    
    .error-details {
        font-size: 0.9rem;
        color: #721c24;
        background-color: #f5c6cb;
        padding: 10px;
        border-radius: 4px;
        margin-top: 10px;
        border-left: 4px solid #dc3545;
        display: none;
    }
    
    .error-toggle {
        color: #0056b3;
        cursor: pointer;
        font-size: 0.8rem;
        margin-top: 5px;
        text-decoration: underline;
    }
    
    .error-toggle:hover {
        color: #003d82;
    }
    
    .warning-container {
        background-color: #fff3cd;
        border: 2px solid #ffeaa7;
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
        color: #856404;
    }
    
    /* Radio Button Styling */
    .radio-group-custom {
        margin: 15px 0;
        padding: 10px;
        border-radius: 8px;
        background-color: #f8f9fa;
    }
    
    .radio-label-custom {
        display: block;
        position: relative;
        padding-left: 35px;
        margin-bottom: 12px;
        cursor: pointer;
        font-size: 1rem;
        user-select: none;
        color: #2c3e50;
        transition: all 0.3s;
    }
    
    .radio-label-custom:hover {
        color: #3498db;
    }
    
    .radio-label-custom input {
        position: absolute;
        opacity: 0;
        cursor: pointer;
    }
    
    .checkmark-custom {
        position: absolute;
        top: 0;
        left: 0;
        height: 22px;
        width: 22px;
        background-color: #e9ecef;
        border-radius: 50%;
        transition: all 0.3s;
        border: 2px solid #dee2e6;
    }
    
    .radio-label-custom:hover input ~ .checkmark-custom {
        background-color: #ddd;
        border-color: #3498db;
    }
    
    .radio-label-custom input:checked ~ .checkmark-custom {
        background-color: #3498db;
        border-color: #2980b9;
    }
    
    .checkmark-custom:after {
        content: "";
        position: absolute;
        display: none;
    }
    
    .radio-label-custom input:checked ~ .checkmark-custom:after {
        display: block;
    }
    
    .radio-label-custom .checkmark-custom:after {
        top: 5px;
        left: 5px;
        width: 8px;
        height: 8px;
        border-radius: 50%;
        background: white;
    }
    
    /* Card Style Radio Buttons */
    .card-radio-container {
        display: flex;
        flex-wrap: wrap;
        gap: 10px;
        margin: 15px 0;
    }
    
    .card-radio-item {
        flex: 1;
        min-width: 120px;
    }
    
    .card-radio-input {
        display: none;
    }
    
    .card-radio-label {
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: 15px 10px;
        background: #f8f9fa;
        border: 2px solid #e9ecef;
        border-radius: 8px;
        text-align: center;
        cursor: pointer;
        transition: all 0.3s;
        height: 100%;
    }
    
    .card-radio-label:hover {
        border-color: #3498db;
        background: #e3f2fd;
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    
    .card-radio-input:checked + .card-radio-label {
        background: linear-gradient(135deg, #3498db, #2980b9);
        color: white;
        border-color: #2980b9;
        box-shadow: 0 6px 12px rgba(52, 152, 219, 0.3);
    }
    
    .card-radio-label i {
        font-size: 1.5rem;
        margin-bottom: 8px;
    }
    
    .card-radio-label small {
        font-size: 0.8rem;
        opacity: 0.8;
        margin-top: 5px;
    }
    
    /* Style for the latest response box */
    .latest-response-box {
        background-color: #f8f9fa;
        border: 2px solid #1976d2;
        border-radius: 10px;
        padding: 20px;
        margin: 15px 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    
    .latest-response-title {
        color: #1976d2;
        font-weight: bold;
        font-size: 1.2rem;
        margin-bottom: 10px;
        border-bottom: 2px solid #e3f2fd;
        padding-bottom: 8px;
    }
    
    .response-content {
        font-size: 1rem;
        line-height: 1.6;
        color: #333;
        max-height: 400px;
        overflow-y: auto;
        padding: 10px;
        background-color: white;
        border-radius: 5px;
        border: 1px solid #e0e0e0;
    }
    
    /* Style for chat history messages */
    .chat-history-message {
        padding: 12px;
        border-radius: 8px;
        margin: 8px 0;
        background-color: #f5f5f5;
        border-left: 4px solid #1976d2;
    }
    
    .chat-history-user {
        background-color: #e3f2fd;
        border-left-color: #2196f3;
    }
    
    .chat-history-assistant {
        background-color: #f1f8e9;
        border-left-color: #4caf50;
    }
    
    /* Style for timestamp */
    .timestamp {
        font-size: 0.8rem;
        color: #666;
        font-style: italic;
        margin-top: 5px;
    }
    
    /* Style for suggestions */
    .suggestion-button {
        margin: 5px 0;
    }
    
    /* Style pills for settings */
    .style-pill {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-size: 0.85rem;
        margin-right: 0.5rem;
        margin-bottom: 0.5rem;
        background-color: #4caf50;
        color: white;
        font-weight: bold;
    }
    
    /* Enhanced document display */
    .document-section {
        background-color: #fff;
        border: 1px solid #ddd;
        border-radius: 5px;
        padding: 15px;
        margin: 10px 0;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    
    .document-header {
        background-color: #f8f9fa;
        padding: 10px;
        border-left: 4px solid #1976d2;
        margin-bottom: 15px;
    }
    
    .document-meta {
        font-size: 0.85rem;
        color: #666;
        display: flex;
        gap: 15px;
        margin-bottom: 10px;
    }
    
    .download-buttons {
        display: flex;
        gap: 10px;
        flex-wrap: wrap;
        margin: 15px 0;
    }
    
    .diagram-placeholder {
        background-color: #f0f8ff;
        border: 1px dashed #4caf50;
        padding: 15px;
        border-radius: 5px;
        margin: 15px 0;
        text-align: center;
    }
    
    .diagram-placeholder code {
        background-color: #fff;
        padding: 10px;
        border-radius: 3px;
        display: block;
        text-align: left;
        margin: 10px 0;
        font-family: monospace;
        white-space: pre-wrap;
    }
    
    .comprehensive-badge {
        background-color: #ff9800;
        color: white;
        padding: 2px 8px;
        border-radius: 12px;
        font-size: 0.75rem;
        font-weight: bold;
        margin-left: 8px;
    }
    
    /* Form Group Styling */
    .form-group-custom {
        margin-bottom: 20px;
        padding: 15px;
        background: white;
        border-radius: 8px;
        border: 1px solid #e0e0e0;
    }
    
    .form-label-custom {
        display: block;
        margin-bottom: 8px;
        font-weight: 600;
        color: #2c3e50;
        font-size: 1.1rem;
    }
    
    .form-label-custom.required::after {
        content: " *";
        color: #e74c3c;
    }
    
    /* Required Field Indicator */
    .required-field::after {
        content: " *";
        color: #e74c3c;
    }
    
    /* Validation Styles */
    .validation-error {
        color: #e74c3c;
        font-size: 0.9rem;
        margin-top: 5px;
        padding: 8px;
        background: #fde8e8;
        border-radius: 4px;
        border-left: 4px solid #e74c3c;
    }
    
    /* Success Message */
    .success-message {
        background: linear-gradient(135deg, #2ecc71, #27ae60);
        color: white;
        padding: 15px 20px;
        border-radius: 8px;
        margin: 20px 0;
        display: none;
        animation: slideIn 0.5s ease;
    }
    
    @keyframes slideIn {
        from {
            opacity: 0;
            transform: translateY(-10px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    /* Document Selection Grid */
    .doc-grid {
        display: grid;
        grid-template-columns: repeat(auto-fill, minmax(180px, 1fr));
        gap: 15px;
        margin: 20px 0;
    }
    
    .doc-card {
        background: white;
        border: 2px solid #e0e0e0;
        border-radius: 10px;
        padding: 20px;
        text-align: center;
        cursor: pointer;
        transition: all 0.3s;
        position: relative;
        overflow: hidden;
    }
    
    .doc-card:hover {
        border-color: #3498db;
        transform: translateY(-5px);
        box-shadow: 0 8px 16px rgba(0,0,0,0.1);
    }
    
    .doc-card.selected {
        border-color: #3498db;
        background: linear-gradient(135deg, #e3f2fd, #bbdefb);
        box-shadow: 0 6px 12px rgba(52, 152, 219, 0.2);
    }
    
    .doc-card i {
        font-size: 2rem;
        color: #3498db;
        margin-bottom: 10px;
    }
    
    .doc-card.selected i {
        color: #1976d2;
    }
    
    .doc-title {
        font-weight: 600;
        color: #2c3e50;
        font-size: 0.9rem;
        margin-bottom: 5px;
    }
    
    .doc-desc {
        font-size: 0.75rem;
        color: #7f8c8d;
        margin-top: 8px;
    }
    
    /* Generate Button */
    .generate-btn {
        background: linear-gradient(135deg, #3498db, #2980b9);
        color: white;
        border: none;
        padding: 15px 30px;
        border-radius: 8px;
        font-size: 1.1rem;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.3s;
        width: 100%;
        margin-top: 20px;
    }
    
    .generate-btn:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 16px rgba(52, 152, 219, 0.3);
    }
    
    .generate-btn:active {
        transform: translateY(0);
    }
    
    /* Responsive */
    @media (max-width: 768px) {
        .doc-grid {
            grid-template-columns: repeat(auto-fill, minmax(150px, 1fr));
        }
        
        .card-radio-item {
            min-width: 100px;
        }
    }
</style>

<script>
function toggleErrorDetails(id) {
    var element = document.getElementById(id);
    if (element.style.display === "none" || element.style.display === "") {
        element.style.display = "block";
    } else {
        element.style.display = "none";
    }
}
</script>
""", unsafe_allow_html=True)

# -------------------------------------------------------------------------
# MAPPING CONFIGURATIONS FOR STATIC DOCUMENT SERVING
# -------------------------------------------------------------------------

# Maps UI Selection strings to Actual File Names in folders
DOC_FILE_MAPPING = {
    "Business Requirements Document (BRD)": "Business Requirements Document.docx",
    "Functional Requirements Document (FRD)": "Functional Requirements Document.docx",
    "Non-Functional Requirements (NFR)": "NFR.docx",
    "User Stories & Acceptance Criteria": "userstories.docx",
    "Stakeholder Analysis Matrix": "Stakeholder Analysis Report.docx",
    "SWOT Analysis Report": "SWOT Analysis Report.docx",
    "Process Flow Diagrams (BPMN)": "Process Flow Documentation.docx"
}

# Helper to map Project Name to Folder Name
def get_folder_name_for_project(project_name):
    """Maps a project selector name to the file system folder name."""
    name_lower = project_name.lower()
    
    if "buddy" in name_lower:
        return "Buddy"
    elif "grader" in name_lower:
        return "Grader"
    elif "hezop" in name_lower:
        return "Hezop"
    elif "price" in name_lower or "prediction" in name_lower:
        return "Price-pred"
    
    # Default fallback if no match found
    return "Buddy" 

# Global error display function
def display_error(error_msg, detailed_error=None, error_type="error"):
    """Display error messages without blocking the app"""
    try:
        error_id = f"error_{datetime.now().strftime('%H%M%S%f')}"
        
        if error_type == "error":
            st.markdown(f"""
            <div class="error-container">
                <div class="error-title">⚠️ Error: {error_msg}</div>
                <div class="error-toggle" onclick="toggleErrorDetails('{error_id}')">
                    Click here to view technical details
                </div>
                <div id="{error_id}" class="error-details">
                    {detailed_error or "No additional details available"}
                </div>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
            <div class="warning-container">
                <div class="error-title">⚠️ Warning: {error_msg}</div>
                {f'<div style="font-size: 0.9rem; margin-top: 5px;">{detailed_error}</div>' if detailed_error else ''}
            </div>
            """, unsafe_allow_html=True)
    except Exception as e:
        # Fallback to simple error display
        st.error(f"Error: {error_msg}")

# Safe initialization function
def safe_initialize():
    """Initialize session state with error handling"""
    try:
        if 'agent' not in st.session_state:
            st.session_state.agent = BusinessAnalystAgent()
        
        if 'project_selector' not in st.session_state:
            st.session_state.project_selector = ProjectSelector()
        
        if 'selected_project' not in st.session_state:
            st.session_state.selected_project = None
        
        if 'current_project_name' not in st.session_state:
            st.session_state.current_project_name = ""
        
        # Chat settings
        if 'response_style' not in st.session_state:
            st.session_state.response_style = "detailed"
        
        if 'response_scope' not in st.session_state:
            st.session_state.response_scope = "specific"
        
        # Chat input
        if 'user_input' not in st.session_state:
            st.session_state.user_input = ""
        
        # Auto-send flag for suggestions
        if 'auto_send' not in st.session_state:
            st.session_state.auto_send = False
        
        # Show/hide chat history
        if 'show_chat_history' not in st.session_state:
            st.session_state.show_chat_history = False
        
        # Store messages for current session
        if 'session_messages' not in st.session_state:
            st.session_state.session_messages = {}
        
        # Store latest response
        if 'latest_response' not in st.session_state:
            st.session_state.latest_response = ""
        
        # Store generated documents (Dictionary for binary data now)
        if 'generated_documents' not in st.session_state:
            st.session_state.generated_documents = {}
        
        # Store last document for download
        if 'last_document' not in st.session_state:
            st.session_state.last_document = None
        if 'last_document_name' not in st.session_state:
            st.session_state.last_document_name = ""
        
        # Document generation progress
        if 'doc_generation_progress' not in st.session_state:
            st.session_state.doc_generation_progress = 0
        
        # Document selection state
        if 'selected_document_type' not in st.session_state:
            st.session_state.selected_document_type = None
        
        if 'document_generation_form' not in st.session_state:
            st.session_state.document_generation_form = {}
        
        # Initialize document form values
        if 'doc_detail_level' not in st.session_state:
            st.session_state.doc_detail_level = "Comprehensive"
        
        if 'doc_format_tables' not in st.session_state:
            st.session_state.doc_format_tables = True
        
        if 'doc_validation_errors' not in st.session_state:
            st.session_state.doc_validation_errors = {}
            
        if 'app_errors' not in st.session_state:
            st.session_state.app_errors = []
            
    except Exception as e:
        display_error("Failed to initialize session state", str(e))
        st.session_state.selected_project = None
        st.session_state.session_messages = {}

# Function to clean markdown formatting
def clean_markdown_formatting(text):
    """Remove markdown formatting from text"""
    try:
        if not text:
            return ""
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            if line.startswith('# '): line = line[2:].strip()
            elif line.startswith('## '): line = line[3:].strip()
            elif line.startswith('### '): line = line[4:].strip()
            elif line.startswith('#### '): line = line[5:].strip()
            
            line = line.replace('**', '').replace('*', '')
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    except Exception as e:
        display_error("Error cleaning markdown formatting", str(e))
        return text or ""

# Function to clean response text for display
def clean_response_text(text):
    """Remove unwanted formatting and extract clean text"""
    try:
        if not text:
            return ""
        text = clean_markdown_formatting(text)
        import re
        text = re.sub(r'<[^>]+>', '', text)
        return text
    except Exception as e:
        display_error("Error cleaning response text", str(e))
        return text or ""

# Function to create a Word document from text (Used only for Chat responses now)
def create_comprehensive_word_document(content, filename, project):
    """Create a comprehensive Word document with professional formatting"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        
        doc = Document()
        styles = doc.styles
        
        # Title style
        try:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Calibri Light'
            title_style.font.size = Pt(28)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 0, 0)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(24)
        except: pass 
        
        # Cover Page
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"{project.get('product_name', 'Project')}\n")
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 32, 96)
        
        doc.add_paragraph().add_run("\n")
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Chat Analysis Report\n")
        run.font.size = Pt(24)
        run.font.bold = True
        
        doc.add_paragraph().add_run("\n\n")
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n")
        
        doc.add_page_break()
        
        # Add content
        doc.add_paragraph(clean_markdown_formatting(content))
        
        import io
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream
        
    except Exception as e:
        display_error("Error creating Word document", str(e))
        return None

# Function to get or create session messages for a project
def get_session_messages(project_name):
    try:
        if project_name not in st.session_state.session_messages:
            st.session_state.session_messages[project_name] = []
        return st.session_state.session_messages[project_name]
    except Exception as e:
        display_error("Error getting session messages", str(e))
        return []

# Function to add message to session
def add_to_session_messages(project_name, role, content, timestamp=None):
    try:
        messages = get_session_messages(project_name)
        messages.append({
            "role": role,
            "content": clean_response_text(content),
            "timestamp": timestamp or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
        st.session_state.session_messages[project_name] = messages[-20:]
    except Exception as e:
        display_error("Error adding message to session", str(e))

# Safe function to get projects
def safe_get_projects():
    try:
        return st.session_state.project_selector.get_all_projects()
    except Exception as e:
        display_error("Could not load projects", str(e))
        return []

# Safe function to analyze with options (Kept for Chat)
def safe_analyze_with_options(project, user_input, response_style="detailed", scope="specific"):
    try:
        return st.session_state.agent.analyze_with_options(
            project,
            user_input,
            response_style=response_style,
            scope=scope
        )
    except Exception as e:
        return f"Error analyzing request: {str(e)}"

# -------------------------------------------------------------------------
# NEW FUNCTION FOR FILE-BASED GENERATION (NO API)
# -------------------------------------------------------------------------
def safe_generate_document_from_file(project, document_type):
    """
    Loads a pre-existing Word document from the file system based on project and doc type.
    Replaces the AI generation logic to save API costs.
    """
    try:
        # 1. Determine Folder Name
        project_name = project.get('product_name', '')
        folder_name = get_folder_name_for_project(project_name)
        
        # 2. Determine File Name
        file_name = DOC_FILE_MAPPING.get(document_type)
        if not file_name:
            raise ValueError(f"No file mapping found for document type: {document_type}")
            
        # 3. Construct Path
        # Assumes 'Documents' folder is in the same directory as this script
        file_path = os.path.join("Documents", folder_name, file_name)
        
        # 4. Simulate Processing Delay
        time.sleep(1) 
        
        # 5. Read File
        if os.path.exists(file_path):
            with open(file_path, "rb") as f:
                file_bytes = f.read()
            return file_bytes, file_name
        else:
            # Fallback for error handling
            error_msg = f"File not found: Documents/{folder_name}/{file_name}"
            display_error(error_msg, "Please ensure the Documents folder structure is correct.", "warning")
            return None, None
            
    except Exception as e:
        display_error(f"Error loading document", str(e))
        return None, None

def extract_text_preview_from_bytes(file_bytes):
    """Extracts text from docx bytes for preview in UI."""
    try:
        if not file_bytes: return ""
        import io
        from docx import Document
        
        source_stream = io.BytesIO(file_bytes)
        document = Document(source_stream)
        
        full_text = []
        for para in document.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        return "\n\n".join(full_text)
    except ImportError:
        return "Preview unavailable (python-docx not installed). Please download the file to view."
    except Exception as e:
        return f"Preview unavailable: {str(e)}"


# Header with error boundary
try:
    st.title("🤖 Business Analyst Agent")
    st.markdown("AI-powered business analysis for your projects")
except Exception as e:
    display_error("Error displaying header", str(e))

# Initialize session state safely
try:
    safe_initialize()
except Exception as e:
    display_error("Critical initialization error", str(e))
    st.stop()

# Sidebar with error boundary
try:
    with st.sidebar:
        st.markdown("## 📋 Project Management")
        
        # Get all projects safely
        projects = safe_get_projects()
        project_names = ["-- Select a Project --"]
        if projects:
            project_names.extend([p.get("product_name", "Unknown") for p in projects])
        
        # Project selection
        selected_project_name = st.selectbox(
            "Select a project:",
            project_names,
            index=0
        )
        
        # Handle project selection
        if selected_project_name != "-- Select a Project --":
            try:
                if selected_project_name != st.session_state.current_project_name:
                    st.session_state.selected_project = st.session_state.project_selector.get_project_by_name(selected_project_name)
                    st.session_state.current_project_name = selected_project_name
                    st.session_state.user_input = ""
                    st.rerun()
            except Exception as e:
                display_error("Error selecting project", str(e))
        
        if st.session_state.selected_project:
            try:
                project = st.session_state.selected_project
                
                st.markdown(f"### 🎯 {project['product_name']}")
                st.caption(f"Source Folder: {get_folder_name_for_project(project['product_name'])}")
                
                # Get actual chat history from agent
                try:
                    actual_history = st.session_state.agent.get_project_history(project['product_name'])
                    st.metric("Total Messages", len(actual_history))
                except:
                    st.metric("Total Messages", 0)
                
                st.markdown("---")
                
                # Response Settings
                st.markdown("### ⚙️ Response Settings")
                
                col1, col2 = st.columns(2)
                with col1:
                    style = st.radio(
                        "Style:",
                        ["Detailed", "Simple"],
                        index=0 if st.session_state.response_style == "detailed" else 1,
                        key="style_radio"
                    )
                    st.session_state.response_style = style.lower()
                
                with col2:
                    scope = st.radio(
                        "Scope:",
                        ["Specific", "General"],
                        index=0 if st.session_state.response_scope == "specific" else 1,
                        key="scope_radio"
                    )
                    st.session_state.response_scope = scope.lower()
                
                st.markdown("---")
                
                # Chat Management
                st.markdown("### 💬 Chat Management")
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("🗑️ Clear History", use_container_width=True):
                        try:
                            # Clear both agent history and session messages
                            st.session_state.agent.clear_project_history(project['product_name'])
                            if project['product_name'] in st.session_state.session_messages:
                                st.session_state.session_messages[project['product_name']] = []
                            st.session_state.latest_response = ""
                            st.success(f"Chat history cleared for {project['product_name']}")
                            st.rerun()
                        except Exception as e:
                            display_error("Error clearing history", str(e))
                
                with col2:
                    # Toggle chat history visibility
                    toggle_text = "📜 Show History" if not st.session_state.show_chat_history else "📜 Hide History"
                    if st.button(toggle_text, use_container_width=True):
                        st.session_state.show_chat_history = not st.session_state.show_chat_history
                        st.rerun()
                        
            except Exception as e:
                display_error("Error displaying project details", str(e))
except Exception as e:
    display_error("Sidebar error", str(e))

# Main Content Area
try:
    if not st.session_state.selected_project:
        # Welcome screen
        st.info("👈 Select a project from the sidebar to begin analysis")
        
        # Show available projects
        st.markdown("### Available Projects")
        if projects:
            for project in projects:
                try:
                    with st.expander(f"📁 {project.get('product_name', 'Unknown Project')}"):
                        sections = project.get('sections', {})
                        one_line = sections.get('One-Line Summary', sections.get('Solution', ''))
                        if isinstance(one_line, str) and len(one_line) > 200:
                            one_line = one_line[:200] + "..."
                        st.write(one_line)
                        if st.button(f"Select {project.get('product_name', 'Project')}", key=f"select_{project.get('product_name', '')}"):
                            try:
                                st.session_state.selected_project = project
                                st.session_state.current_project_name = project.get('product_name', '')
                                st.rerun()
                            except Exception as e:
                                display_error("Error selecting project", str(e))
                except Exception as e:
                    display_error(f"Error displaying project {project.get('product_name', 'Unknown')}", str(e))
        else:
            st.warning("No projects available. Check if project data is loaded correctly.")
    
    else:
        # Project is selected
        try:
            project = st.session_state.selected_project
            
            # Project header
            col1, col2, col3 = st.columns([3, 2, 1])
            with col1:
                st.markdown(f"### 💬 Chat: **{project.get('product_name', 'Unknown Project')}**")
            with col2:
                # Style indicators
                style_pill = "🟢 Detailed" if st.session_state.response_style == "detailed" else "⚪ Simple"
                scope_pill = "🎯 Specific" if st.session_state.response_scope == "specific" else "🌐 General"
                st.markdown(
                    f"<div style='display: flex; gap: 0.5rem; flex-wrap: wrap;'>"
                    f"<span class='style-pill active'>{style_pill}</span>"
                    f"<span class='style-pill active'>{scope_pill}</span>"
                    f"</div>",
                    unsafe_allow_html=True
                )
            with col3:
                if st.button("🔄 Refresh"):
                    st.rerun()
            
            # Create tabs
            try:
                tab1, tab2, tab3 = st.tabs(["💬 Chat", "📄 Document Generation", "📊 Project Details"])
                
                # --- TAB 1: CHAT ---
                with tab1:
                    try:
                        # Get session messages for this project
                        session_messages = get_session_messages(project.get('product_name', ''))
                        
                        # --- CHAT INPUT ---
                        st.markdown("### 💬 Your Question")
                        
                        user_input = st.text_area(
                            "Type your question:",
                            value=st.session_state.user_input,
                            placeholder=f"Ask about {project.get('product_name', 'the project')}...",
                            height=80,
                            key="chat_input",
                            label_visibility="collapsed"
                        )
                        
                        # Send button
                        col1, col2 = st.columns([4, 1])
                        with col2:
                            send_button = st.button("🚀 Send", use_container_width=True, type="primary")
                        
                        # Handle send button
                        if send_button and user_input.strip():
                            try:
                                with st.spinner("🤔 Analyzing..."):
                                    # Add user message to session
                                    add_to_session_messages(project.get('product_name', ''), "user", user_input)
                                    
                                    # Get analysis from agent
                                    response = safe_analyze_with_options(
                                        project,
                                        user_input,
                                        response_style=st.session_state.response_style,
                                        scope=st.session_state.response_scope
                                    )
                                    
                                    # Add assistant response to session
                                    add_to_session_messages(project.get('product_name', ''), "assistant", response)
                                    
                                    # Store latest response
                                    st.session_state.latest_response = clean_response_text(response)
                                    
                                    # Clear input
                                    st.session_state.user_input = ""
                                    st.rerun()
                            except Exception as e:
                                display_error("Error processing your question", str(e))
                        
                        # Divider
                        st.markdown("---")
                        
                        # --- LATEST RESPONSE ---
                        if st.session_state.latest_response or session_messages:
                            st.markdown("### 📋 **Latest Response**")
                            
                            latest_response = st.session_state.latest_response
                            if not latest_response and session_messages:
                                assistant_messages = [m for m in session_messages if m.get("role") == "assistant"]
                                if assistant_messages:
                                    latest_response = assistant_messages[-1].get("content", "")
                            
                            if latest_response:
                                st.markdown(f"""
                                <div class="latest-response-box">
                                    <div class="latest-response-title">📋 Business Analyst Response</div>
                                    <div class="response-content">
                                        <strong>{latest_response[:500]}</strong>
                                        {latest_response[500:] if len(latest_response) > 500 else ''}
                                    </div>
                                </div>
                                """, unsafe_allow_html=True)
                        
                        # --- CHAT HISTORY ---
                        if st.session_state.show_chat_history and len(session_messages) > 0:
                            st.markdown("---")
                            st.markdown("### 📜 Chat History")
                            
                            history_messages = session_messages[:-1] if len(session_messages) > 1 else []
                            if history_messages:
                                history_messages_reversed = list(reversed(history_messages))
                                for msg in history_messages_reversed:
                                    role_label = "👤 You" if msg.get("role") == "user" else "🤖 Business Analyst"
                                    css_class = "chat-history-user" if msg.get("role") == "user" else "chat-history-assistant"
                                    
                                    st.markdown(f"""
                                    <div class="chat-history-message {css_class}">
                                        <strong>{role_label}:</strong> {msg.get('content', '')[:150]}...
                                        <div class="timestamp">{msg.get('timestamp', '')}</div>
                                    </div>
                                    """, unsafe_allow_html=True)

                    except Exception as e:
                        display_error("Error in chat tab", str(e))
                
                # --- TAB 2: DOCUMENT GENERATION (MODIFIED TO USE LOCAL FILES) ---
                with tab2:
                    try:
                        st.markdown("## 📄 Document Repository")
                        st.info("Select a document type to retrieve pre-generated professional documents for this project.")
                        
                        # Document type selection
                        st.markdown("### 📋 Select Document Type")
                        
                        document_types = [
                            {"type": "Business Requirements Document (BRD)", "icon": "📋"},
                            {"type": "Functional Requirements Document (FRD)", "icon": "⚙️"},
                            {"type": "Non-Functional Requirements (NFR)", "icon": "📊"},
                            {"type": "User Stories & Acceptance Criteria", "icon": "👥"},
                            {"type": "Stakeholder Analysis Matrix", "icon": "🤝"},
                            {"type": "SWOT Analysis Report", "icon": "🔍"},
                            {"type": "Process Flow Diagrams (BPMN)", "icon": "🔄"},
                        ]
                        
                        # Create document selection grid
                        cols = st.columns(5)
                        for idx, doc_info in enumerate(document_types):
                            col_idx = idx % 5
                            with cols[col_idx]:
                                is_selected = st.session_state.selected_document_type == doc_info["type"]
                                
                                st.markdown(f"""
                                <div class="doc-card {'selected' if is_selected else ''}" onclick="document.getElementById('doc_{idx}').click()">
                                    <div style="font-size: 2rem; margin-bottom: 10px;">{doc_info['icon']}</div>
                                    <div class="doc-title">{doc_info['type'].split('(')[0].strip()}</div>
                                </div>
                                """, unsafe_allow_html=True)
                                
                                if st.button(f"Select", 
                                            key=f"doc_{idx}",
                                            type="primary" if is_selected else "secondary",
                                            use_container_width=True):
                                    st.session_state.selected_document_type = doc_info["type"]
                                    st.rerun()
                        
                        # Show selected document
                        if st.session_state.selected_document_type:
                            st.markdown(f"""
                            <div class="form-group-custom">
                                <div class="form-label-custom">✅ Selected Document</div>
                                <div style="font-size: 1.2rem; color: #2c3e50; padding: 10px; background: #e8f4fc; border-radius: 5px;">
                                    📄 <strong>{st.session_state.selected_document_type}</strong>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                        
                        # Generate Button (Now "Retrieve")
                        st.markdown("---")
                        if st.button("🚀 Generate Document", type="primary", use_container_width=True):
                            if not st.session_state.selected_document_type:
                                display_error("Please select a document type first!", error_type="warning")
                            else:
                                with st.spinner(f"📂 Accessing {st.session_state.selected_document_type}..."):
                                    
                                    # USE LOCAL FILE LOAD LOGIC INSTEAD OF AI AGENT
                                    doc_bytes, file_name = safe_generate_document_from_file(
                                        project=project,
                                        document_type=st.session_state.selected_document_type
                                    )
                                    
                                    if doc_bytes:
                                        # Store for download (Binary)
                                        st.session_state.last_document = doc_bytes
                                        st.session_state.last_document_name = file_name
                                        
                                        # Extract text for preview
                                        preview_text = extract_text_preview_from_bytes(doc_bytes)
                                        
                                        # Store in generated documents mapping
                                        project_name = project.get('product_name', '')
                                        if project_name not in st.session_state.generated_documents:
                                            st.session_state.generated_documents[project_name] = {}
                                        
                                        st.session_state.generated_documents[project_name][st.session_state.selected_document_type] = {
                                            "bytes": doc_bytes,
                                            "preview": preview_text,
                                            "filename": file_name
                                        }
                                        
                                        st.success(f"✅ {st.session_state.selected_document_type} retrieved successfully!")
                        
                        # Display generated documents for this project
                        project_name = project.get('product_name', '')
                        if project_name in st.session_state.generated_documents:
                            st.markdown("---")
                            st.markdown("### 📂 Retrieved Documents")
                            
                            for doc_type, doc_data in st.session_state.generated_documents[project_name].items():
                                try:
                                    with st.expander(f"📄 {doc_type}", expanded=True):
                                        
                                        # Show preview
                                        st.text_area(
                                            f"Preview of {doc_type}",
                                            value=doc_data["preview"][:1500] + "\n...",
                                            height=200,
                                            disabled=True,
                                            key=f"preview_{doc_type}"
                                        )
                                        
                                        # Download button
                                        st.download_button(
                                            label="📄 Download Word File",
                                            data=doc_data["bytes"],
                                            file_name=doc_data["filename"],
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            use_container_width=True,
                                            key=f"dl_{doc_type}",
                                            type="primary"
                                        )
                                        
                                except Exception as e:
                                    display_error(f"Error displaying document {doc_type}", str(e))

                    except Exception as e:
                        display_error("Error in document generation tab", str(e))
                
                # --- TAB 3: PROJECT DETAILS ---
                with tab3:
                    try:
                        st.markdown("## 📊 Project Details")
                        sections = project.get('sections', {})
                        if sections:
                            for section_name, section_content in sections.items():
                                with st.expander(f"**{section_name}**", expanded=(section_name in ["Problem Statement", "Solution"])):
                                    st.write(section_content)
                        else:
                            st.info("No project details available.")
                    except Exception as e:
                        display_error("Error in project details tab", str(e))
                        
            except Exception as e:
                display_error("Error creating tabs", str(e))
                
        except Exception as e:
            display_error("Error displaying project content", str(e))
            st.info("Please select a different project or refresh the page.")

except Exception as e:
    display_error("Unexpected error in main application", str(e))
    if st.button("🔄 Reset Application"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

# Footer
try:
    st.markdown("---")
    st.caption(f"Business Analyst Agent v1.0 • Last updated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
except:
    pass